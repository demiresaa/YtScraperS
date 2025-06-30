import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from datetime import datetime
import numpy as np
import ollama 
import time 
import re 

class YouTubeAnalysisReporter:
    def __init__(self, csv_file_path, keyword, category_definitions,
                 llm_model_name="gemma3:12b", llm_temperature=0.5, llm_max_tokens=4096, llm_context_window=8192):
        """
        Initializes the YouTubeAnalysisReporter.

        Args:
            csv_file_path (str): Path to the CSV file containing comment data.
            keyword (str): The main keyword for the analysis.
            category_definitions (list): A list of dictionaries, where each dictionary
                                         defines a category.
                                         Example: [{"csv_col": "fiyat", "display_name": "Fiyat Değerlendirmeleri"}, ...]
            llm_model_name (str): Name of the Ollama model to use.
            llm_temperature (float): Temperature for LLM generation.
            llm_max_tokens (int): Max tokens for LLM generation.
            llm_context_window (int): Context window size for the LLM.
        """
        self.csv_file_path = csv_file_path
        self.keyword = keyword
        self.category_definitions = category_definitions
        self.llm_model_name = llm_model_name
        self.llm_temperature = llm_temperature
        self.llm_max_tokens = llm_max_tokens
        self.llm_context_window = llm_context_window

        self.df_comments = self._load_and_preprocess_data()

    def _load_and_preprocess_data(self):
        """Loads data from the CSV file and performs initial preprocessing."""
        try:
            df = pd.read_csv(self.csv_file_path)
            print(f"'{self.csv_file_path}' başarıyla okundu. Toplam {len(df)} yorum bulundu.")

            required_cols = ['duygu_tahmini', 'duygu_skoru', 'hesaplanan_tarih']
            missing_cols = [col for col in required_cols if col not in df.columns]
            if missing_cols:
                raise ValueError(f"Gerekli sütunlar CSV dosyasında eksik: {', '.join(missing_cols)}")

            df['duygu_tahmini'] = df['duygu_tahmini'].fillna('neutral').astype(str).str.lower()
            df['duygu_skoru'] = pd.to_numeric(df['duygu_skoru'], errors='coerce').fillna(0)

            if 'hesaplanan_tarih' in df.columns:
                df['hesaplanan_tarih'] = pd.to_datetime(df['hesaplanan_tarih'], errors='coerce')
                df.dropna(subset=['hesaplanan_tarih'], inplace=True)
                if df.empty and not df['hesaplanan_tarih'].isna().all(): # Check if empty after dropna
                     print("UYARI: Geçerli 'hesaplanan_tarih' verisi kalmadı. Zaman serisi analizleri yapılamayacak.")
            else:
                print("UYARI: 'hesaplanan_tarih' sütunu DataFrame'de bulunamadı. Zaman serisi analizleri etkilenebilir.")

            for cat_def in self.category_definitions:
                col_name = cat_def['csv_col']
                if col_name in df.columns:
       
                    try:
                        df[col_name] = df[col_name].fillna(0).astype(int)
                    except ValueError:
                        print(f"UYARI: '{col_name}' sütunu integer'a çevrilemedi. 0 olarak ayarlandı.")
                        df[col_name] = 0

                else:
                    print(f"UYARI: Beklenen kategori sütunu '{col_name}' DataFrame'de bulunamadı. Bu sütun 0'larla oluşturulacak.")
                    df[col_name] = 0
            return df

        except FileNotFoundError:
            print(f"HATA: '{self.csv_file_path}' dosyası bulunamadı. Lütfen geçerli bir dosya yolu sağlayın.")
            raise
        except ValueError as ve:
            print(f"Veri Hatası: {ve}")
            raise
        except Exception as e:
            print(f"Veri yükleme ve ön işleme sırasında beklenmedik hata: {e}")
            raise

    def _ollama_gemma_chat(self, prompt_text):
        """
        Sends a prompt to the Gemma model running on Ollama and gets the response.
        """
        try:
            client = ollama.Client()
            response = client.chat(
                model=self.llm_model_name,
                messages=[
                    {'role': 'system', 'content': 'Sen profesyonel bir veri bilimcisin. Sana verilen talimatlara göre, belirtilen formatta ve ayraçları kullanarak detaylı, objektif ve Türkçe analiz metinleri üreteceksin. Yorumların bilgilendirici ve eyleme geçirilebilir içgörüler sunmalı. Sadece verilen bilgilerle yorum yap. **Markdown veya herhangi bir özel formatlama işareti (örneğin **, ##, *, -, 1., 2.) KULLANMA. Sadece istenen paragraf metnini üret.**'},
                    {'role': 'user', 'content': prompt_text}
                ],
                options={
                    'temperature': self.llm_temperature,
                    'num_ctx': self.llm_context_window,
                    'num_predict': self.llm_max_tokens
                }
            )
            return response['message']['content']
        except Exception as e:
            print(f"LLM ile iletişim hatası: {e}")
            return "LLM_YANIT_HATASI: LLM'den yanıt alınamadı. Lütfen Ollama sunucusunun çalıştığından ve modelin yüklü olduğundan emin olun."

    def _summarize_sentiment_trend(self, df_agg, sentiment_type):
        if df_agg.empty or len(df_agg) < 2:
            return f"Bu dönemde {sentiment_type} duygu skoru için yeterli zaman serisi verisi bulunmamaktadır."

        first_date = df_agg['hesaplanan_tarih'].iloc[0].strftime('%Y-%m-%d')
        last_date = df_agg['hesaplanan_tarih'].iloc[-1].strftime('%Y-%m-%d')
        first_score = df_agg['duygu_skoru'].iloc[0]
        last_score = df_agg['duygu_skoru'].iloc[-1]
        
        overall_change = last_score - first_score
        
        trend_description = ""
        if overall_change > 0.05:
            trend_description = "genel olarak yükseliş eğilimindedir"
        elif overall_change < -0.05:
            trend_description = "genel olarak düşüş eğilimindedir"
        else:
            trend_description = "nispeten sabit bir seyir izlemiştir"

        max_score_val = df_agg['duygu_skoru'].max()
        min_score_val = df_agg['duygu_skoru'].min()
        date_of_max = df_agg.loc[df_agg['duygu_skoru'].idxmax(), 'hesaplanan_tarih'].strftime('%Y-%m-%d')
        date_of_min = df_agg.loc[df_agg['duygu_skoru'].idxmin(), 'hesaplanan_tarih'].strftime('%Y-%m-%d')

        return (f"({first_date} - {last_date} tarihleri arası) Ortalama {sentiment_type} duygu skoru {trend_description}. "
                f"En yüksek skor ({max_score_val:.2f}) {date_of_max} tarihinde, "
                f"en düşük skor ({min_score_val:.2f}) ise {date_of_min} tarihinde kaydedilmiştir.")

    def _sanitize_for_delimiter(self, category_name):
        s = category_name.replace(" ", "_")
        s = re.sub(r'[^\w_]', '', s)
        return s.upper()

    def _parse_llm_section(self, response_text, section_key):
        safe_section_key = re.escape(section_key)
        pattern = rf"\[{safe_section_key}_BASLANGIC\](.*?)\s*\[{safe_section_key}_BITIS\]"
        match = re.search(pattern, response_text, re.DOTALL | re.IGNORECASE)
        if match:
            return match.group(1).strip()
        else:
            error_message = f"LLM_PARSE_HATASI: '{section_key}' bölümü LLM yanıtında bulunamadı veya formatı hatalı."
            print(f"Uyarı: {error_message}")
            return error_message


    def _create_modern_sentiment_pie_chart(self):
        df = self.df_comments
        olumlu_yorum_sayisi = df[df['duygu_tahmini'].str.lower() == 'positive'].shape[0]
        olumsuz_yorum_sayisi = df[df['duygu_tahmini'].str.lower() == 'negative'].shape[0]
        notr_yorum_sayisi = df[df['duygu_tahmini'].str.lower() == 'neutral'].shape[0]
        
        etiketler, boyutlar, renkler = [], [], []
        if olumlu_yorum_sayisi > 0: etiketler.append('Olumlu'); boyutlar.append(olumlu_yorum_sayisi); renkler.append('#2ca02c')
        if olumsuz_yorum_sayisi > 0: etiketler.append('Olumsuz'); boyutlar.append(olumsuz_yorum_sayisi); renkler.append('#d62728')
        if notr_yorum_sayisi > 0: etiketler.append('Nötr'); boyutlar.append(notr_yorum_sayisi); renkler.append('#ff7f0e')

        if sum(boyutlar) == 0:
            plt.figure(figsize=(7, 5)); plt.text(0.5, 0.5, 'Duygu analizi için veri yok.', horizontalalignment='center', verticalalignment='center', fontsize=12, color='gray'); plt.axis('off')
            buffer = BytesIO(); plt.savefig(buffer, format='png', bbox_inches='tight', dpi=150); buffer.seek(0); plt.close(); return buffer

        plt.figure(figsize=(7, 5)); plt.rcParams['font.size'] = 11
        wedges, texts, autotexts = plt.pie(boyutlar,
                explode=[0.05 if etiket == 'Olumsuz' else 0 for etiket in etiketler],
                labels=etiketler, colors=renkler,
                autopct=lambda p: '{:.1f}%\n({:d})'.format(p, int(p/100*sum(boyutlar))) if p > 0 else '',
                shadow=True, startangle=90, pctdistance=0.80, labeldistance=1.1,
                wedgeprops=dict(width=0.4, edgecolor='w'))
        plt.setp(autotexts, size=10, weight="bold", color="white")
        plt.title('Yorumların Genel Duygu Dağılımı', fontsize=16, pad=20)
        plt.axis('equal'); plt.tight_layout()
        buffer = BytesIO(); plt.savefig(buffer, format='png', bbox_inches='tight', dpi=150); buffer.seek(0); plt.close()
        return buffer

    def _create_modern_line_chart(self, df_input, title="Duygu Analizi Zaman Serisi Grafiği"):
        # sourcery skip: extract-method
        if df_input.empty or 'hesaplanan_tarih' not in df_input.columns or \
           'duygu_skoru' not in df_input.columns or 'duygu_tahmini' not in df_input.columns:
            fig, ax = plt.subplots(figsize=(10, 5))
            ax.text(0.5, 0.5, 'Bu grafik için yeterli veri bulunmamaktadır.', horizontalalignment='center', verticalalignment='center', transform=ax.transAxes, fontsize=12, color='gray')
            ax.set_xticks([]); ax.set_yticks([]); ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False); ax.spines['bottom'].set_visible(False); ax.spines['left'].set_visible(False)
            plt.title(title, fontsize=16, color='#333333', pad=20)
            buffer = BytesIO(); plt.savefig(buffer, format='png', bbox_inches='tight', dpi=150); buffer.seek(0); plt.close(fig)
            return buffer

        df_copy = df_input.copy(); df_copy['hesaplanan_tarih'] = pd.to_datetime(df_copy['hesaplanan_tarih'], errors='coerce'); df_copy.dropna(subset=['hesaplanan_tarih'], inplace=True)
        if df_copy.empty:
            fig, ax = plt.subplots(figsize=(10, 5)); ax.text(0.5, 0.5, 'Geçerli tarih formatında veri bulunamadı.', horizontalalignment='center', verticalalignment='center', transform=ax.transAxes, fontsize=12, color='gray')
            ax.set_xticks([]); ax.set_yticks([]); ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False); ax.spines['bottom'].set_visible(False); ax.spines['left'].set_visible(False)
            plt.title(title, fontsize=16, color='#333333', pad=20)
            buffer = BytesIO(); plt.savefig(buffer, format='png', bbox_inches='tight', dpi=150); buffer.seek(0); plt.close(fig)
            return buffer
            
        df_pos = df_copy[df_copy['duygu_tahmini'].str.lower() == 'positive'].copy()
        df_neg = df_copy[df_copy['duygu_tahmini'].str.lower() == 'negative'].copy()
        df_neutral = df_copy[df_copy['duygu_tahmini'].str.lower() == 'neutral'].copy()

        df_pos_agg = df_pos.groupby(pd.Grouper(key='hesaplanan_tarih', freq='D'))['duygu_skoru'].mean().reset_index() if not df_pos.empty else pd.DataFrame(columns=['hesaplanan_tarih', 'duygu_skoru'])
        df_neg_agg = df_neg.groupby(pd.Grouper(key='hesaplanan_tarih', freq='D'))['duygu_skoru'].mean().reset_index() if not df_neg.empty else pd.DataFrame(columns=['hesaplanan_tarih', 'duygu_skoru'])
        df_neutral_agg = df_neutral.groupby(pd.Grouper(key='hesaplanan_tarih', freq='D'))['duygu_skoru'].mean().reset_index() if not df_neutral.empty else pd.DataFrame(columns=['hesaplanan_tarih', 'duygu_skoru'])

        for df_a in [df_pos_agg, df_neg_agg, df_neutral_agg]:
            if not df_a.empty: df_a.sort_values(by='hesaplanan_tarih', inplace=True); df_a.dropna(subset=['duygu_skoru'], inplace=True)

        fig, ax = plt.subplots(figsize=(10, 5)); plt.style.use('seaborn-v0_8-whitegrid'); ax.set_facecolor("white")
        positive_color, negative_color, neutral_color = '#2ca02c', '#d62728', '#ff7f0e'
        line_width, marker_size, alpha = 2.5, 6, 0.8
        plotted_something = False

        if not df_pos_agg.empty and pd.api.types.is_numeric_dtype(df_pos_agg['duygu_skoru']):
            ax.plot(df_pos_agg['hesaplanan_tarih'], df_pos_agg['duygu_skoru'], marker='o', markersize=marker_size, linestyle='-', color=positive_color, linewidth=line_width, alpha=alpha, label='Pozitif Ortalama Skor'); plotted_something = True
        if not df_neg_agg.empty and pd.api.types.is_numeric_dtype(df_neg_agg['duygu_skoru']):
            ax.plot(df_neg_agg['hesaplanan_tarih'], df_neg_agg['duygu_skoru'], marker='X', markersize=marker_size, linestyle='--', color=negative_color, linewidth=line_width, alpha=alpha, label='Negatif Ortalama Skor'); plotted_something = True
        if not df_neutral_agg.empty and pd.api.types.is_numeric_dtype(df_neutral_agg['duygu_skoru']):
            ax.plot(df_neutral_agg['hesaplanan_tarih'], df_neutral_agg['duygu_skoru'], marker='s', markersize=marker_size, linestyle=':', color=neutral_color, linewidth=line_width, alpha=alpha, label='Nötr Ortalama Skor'); plotted_something = True

        if not plotted_something:
            ax.text(0.5, 0.5, 'Bu grafik için çizilecek veri bulunmamaktadır.', horizontalalignment='center', verticalalignment='center', transform=ax.transAxes, fontsize=12, color='gray')
            ax.set_xticks([]); ax.set_yticks([])
        else:
            ax.xaxis.set_major_locator(mdates.AutoDateLocator(minticks=5, maxticks=10)); ax.xaxis.set_major_formatter(mdates.DateFormatter('%Y-%m-%d')); plt.gcf().autofmt_xdate(rotation=30)
            ax.legend(loc='best', frameon=True, fontsize=10, facecolor='white', framealpha=0.7)
            all_scores = pd.concat([df_pos_agg['duygu_skoru'], df_neg_agg['duygu_skoru'], df_neutral_agg['duygu_skoru']])
            if not all_scores.empty and not all_scores.isna().all():
                min_val, max_val = all_scores.min(), all_scores.max()
                if pd.notna(min_val) and pd.notna(max_val):
                    if min_val == max_val: ax.set_ylim(min_val - 0.1 , max_val + 0.1)
                    else: ax.set_ylim(min_val - abs(min_val*0.1) - 0.05 , max_val + abs(max_val*0.1) + 0.05)

        ax.set_xlabel('Tarih', fontsize=12, color='#555555'); ax.set_ylabel('Ortalama Duygu Skoru', fontsize=12, color='#555555')
        ax.set_title(title, fontsize=16, color='#333333', pad=20); ax.grid(True, linestyle=':', alpha=0.7)
        ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False); ax.spines['bottom'].set_color('gray'); ax.spines['left'].set_color('gray')
        plt.tight_layout(); buffer = BytesIO(); plt.savefig(buffer, format='png', bbox_inches='tight', dpi=150); buffer.seek(0); plt.close(fig)
        return buffer

    def _define_document_styles(self, document):
        styles = document.styles
        try:
            title_style = styles['TitleStyle']
        except KeyError:
            title_style = styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.base_style = styles['Heading 1']; font = title_style.font; font.name = 'Arial'; font.size = Pt(24); font.bold = True; font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER; title_style.paragraph_format.space_after = Pt(18)
        try:
            heading1_style = styles['Heading1Style']
        except KeyError:
            heading1_style = styles.add_style('Heading1Style', WD_STYLE_TYPE.PARAGRAPH)
            heading1_style.base_style = styles['Heading 1']; font = heading1_style.font; font.name = 'Calibri Light'; font.size = Pt(16); font.bold = True; font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
            heading1_style.paragraph_format.space_before = Pt(18); heading1_style.paragraph_format.space_after = Pt(10)
            p_border = OxmlElement('w:pBdr'); bottom_border = OxmlElement('w:bottom'); bottom_border.set(qn('w:val'), 'single'); bottom_border.set(qn('w:sz'), '6'); bottom_border.set(qn('w:space'), '1'); bottom_border.set(qn('w:color'), '2E74B5')
            p_border.append(bottom_border); heading1_style.element.pPr.append(p_border)
        try:
            heading2_style = styles['Heading2Style']
        except KeyError:
            heading2_style = styles.add_style('Heading2Style', WD_STYLE_TYPE.PARAGRAPH)
            heading2_style.base_style = styles['Heading 2']; font = heading2_style.font; font.name = 'Calibri Light'; font.size = Pt(13); font.bold = True; font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
            heading2_style.paragraph_format.space_before = Pt(12); heading2_style.paragraph_format.space_after = Pt(6)
        try:
            normal_style = styles['NormalTextStyle']
        except KeyError:
            normal_style = styles.add_style('NormalTextStyle', WD_STYLE_TYPE.PARAGRAPH)
            normal_style.base_style = styles['Normal']; font = normal_style.font; font.name = 'Calibri'; font.size = Pt(11)
            normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; normal_style.paragraph_format.line_spacing = 1.15; normal_style.paragraph_format.space_after = Pt(8)
        try:
            caption_style = styles['CaptionStyle']
        except KeyError:
            caption_style = styles.add_style('CaptionStyle', WD_STYLE_TYPE.PARAGRAPH)
            caption_style.base_style = styles['Caption']; font = caption_style.font; font.name = 'Calibri'; font.size = Pt(9); font.italic = True
            caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER; caption_style.paragraph_format.space_before = Pt(2); caption_style.paragraph_format.space_after = Pt(10)
        
        return title_style, heading1_style, heading2_style, normal_style, caption_style

    def _set_cell_shading(self, cell, fill_color):
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), fill_color)
            shading_elm.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shading_elm)

    def generate_report(self, output_filename="Profesyonel_YouTube_Analiz_Raporu_LLM_Unified.docx"):
        if self.df_comments.empty:
            print("DataFrame boş. Rapor oluşturulamıyor.")
            return

        document = Document()
        title_style, heading1_style, heading2_style, normal_style, caption_style = self._define_document_styles(document)

        total_comments = len(self.df_comments)
        positive_comments_count = self.df_comments[self.df_comments['duygu_tahmini'].str.lower() == 'positive'].shape[0]
        negative_comments_count = self.df_comments[self.df_comments['duygu_tahmini'].str.lower() == 'negative'].shape[0]
        neutral_comments_count = self.df_comments[self.df_comments['duygu_tahmini'].str.lower() == 'neutral'].shape[0]
        video_count_placeholder = self.df_comments['video_id'].nunique() if 'video_id' in self.df_comments else 7 # Placeholder if no video_id

        # Category data preparation using self.category_definitions
        category_counts_data = []
        for cat_def in self.category_definitions:
            csv_col_name = cat_def['csv_col']
            display_name = cat_def['display_name']
            count = self.df_comments[self.df_comments[csv_col_name] == 1].shape[0] if csv_col_name in self.df_comments.columns else 0
            category_counts_data.append({'name': display_name, 'count': count, 'csv_col': csv_col_name})
        
        kategori_adlari_str_yontem = ", ".join([item['name'] for item in category_counts_data[:-1]]) + \
                                     (" ve " + category_counts_data[-1]['name'] if len(category_counts_data) > 1 else \
                                      (category_counts_data[0]['name'] if category_counts_data else "tanımlı kategori yok"))

        category_counts_summary_for_llm = "\n".join([f"- {item['name']}: {item['count']} yorum" for item in category_counts_data])
        
        # Check for 'hesaplanan_tarih' before attempting time series aggregation
        general_pos_trend_summary = "Zaman serisi verisi yok."
        general_neg_trend_summary = "Zaman serisi verisi yok."
        if 'hesaplanan_tarih' in self.df_comments.columns and not self.df_comments['hesaplanan_tarih'].isna().all():
            df_pos_agg_general = self.df_comments[self.df_comments['duygu_tahmini'].str.lower() == 'positive'].groupby(pd.Grouper(key='hesaplanan_tarih', freq='D'))['duygu_skoru'].mean().reset_index()
            df_neg_agg_general = self.df_comments[self.df_comments['duygu_tahmini'].str.lower() == 'negative'].groupby(pd.Grouper(key='hesaplanan_tarih', freq='D'))['duygu_skoru'].mean().reset_index()
            # df_neutral_agg_general = self.df_comments[self.df_comments['duygu_tahmini'].str.lower() == 'neutral'].groupby(pd.Grouper(key='hesaplanan_tarih', freq='D'))['duygu_skoru'].mean().reset_index() # Not used in prompt currently
            general_pos_trend_summary = self._summarize_sentiment_trend(df_pos_agg_general, "pozitif")
            general_neg_trend_summary = self._summarize_sentiment_trend(df_neg_agg_general, "negatif")
        
        # --- TEK BÜYÜK PROMPT OLUŞTURMA ---
        master_prompt_parts = []
        master_prompt_parts.append(
            f"Aşağıdaki talimatlara ve verilere dayanarak bir YouTube yorum analizi raporu için Türkçe metinler üreteceksin. "
            f"Her bölüm için istenen metni, bölümün adını taşıyan özel ayraçlar arasına (örneğin, [BOLUM_ADI_BASLANGIC] metin... [BOLUM_ADI_BITIS]) yaz. "
            f"Yorumların kısa, öz, profesyonel ve sadece verilen bilgilere dayalı olmalıdır. Her bölüm için tek paragraf yeterlidir.\n\n"
            f"--- GENEL VERİLER ---\n"
            f"Anahtar Kelime: '{self.keyword}'\n"
            f"Toplam Yorum Sayısı: {total_comments}\n"
            f"Pozitif Yorum Sayısı: {positive_comments_count} (Yüzde: {(positive_comments_count/total_comments*100) if total_comments > 0 else 0:.1f}%)\n"
            f"Negatif Yorum Sayısı: {negative_comments_count} (Yüzde: {(negative_comments_count/total_comments*100) if total_comments > 0 else 0:.1f}%)\n"
            f"Nötr Yorum Sayısı: {neutral_comments_count} (Yüzde: {(neutral_comments_count/total_comments*100) if total_comments > 0 else 0:.1f}%)\n"
            f"Taranan Video Sayısı (Tahmini): {video_count_placeholder}\n"
            f"Ana Kategoriler: {kategori_adlari_str_yontem}\n"
            f"Kategori Bazlı Yorum Sayıları:\n{category_counts_summary_for_llm}\n"
            f"Genel Pozitif Duygu Trend Özeti: {general_pos_trend_summary}\n"
            f"Genel Negatif Duygu Trend Özeti: {general_neg_trend_summary}\n"
        )
        master_prompt_parts.append(
            "\n\n--- BÖLÜM: ÖZET ---\n"
            "[OZET_BASLANGIC]\n"
            "Yukarıdaki genel verilere dayanarak, bu YouTube yorum analizi projesinin amacını, yöntemini (veri toplama, duygu analizi, kategorizasyon), kullanılan verileri ve temel bulguları içeren, anlaşılır, profesyonel 2 paragraf halinde bir özet yaz. Anahtar kelimeyi belirt. Amaç ve Kapsam GİRİŞTE bir Neden bu analizi yaptık?” ve “Bu raporun kapsadığı dönem/metod nedir?” gibi sorulara cevap verilmeli. Anahtar Bulgular (Key Findings) Rakamların yanı sıra “en kritik 3–5 bulgu” kısa kısa bahsetmeli. \n"
            "[OZET_BITIS]"
        )
        master_prompt_parts.append(
            "\n\n--- BÖLUM: PASTA GRAFİK YORUMU ---\n"
            "[PASTA_YORUM_BASLANGIC]\n"
            "Yukarıdaki genel duygu dağılım istatistiklerine (toplam, pozitif, negatif, nötr yorum sayıları ve yüzdeleri) dayanarak, yorumların genel duygu dağılımı hakkında TEK paragraflık uzun bir yorum yaz. Genel algının ne yönde olduğuna dair bir çıkarım yap.\n"
            "[PASTA_YORUM_BITIS]"
        )
        master_prompt_parts.append(
            "\n\n--- BÖLÜM: GENEL ZAMAN SERİSİ YORUMU ---\n"
            "[GENEL_TREND_YORUM_BASLANGIC]\n"
            "Yukarıda verilen 'Genel Pozitif Duygu Trend Özeti' ve 'Genel Negatif Duygu Trend Özeti' bilgilerine dayanarak, zaman içinde genel duygu skorlarının değişimi hakkında TEK paragraflık uzun bir yorum yaz. Belirgin eğilimleri, zirve/dip noktalarını (verildiyse) vurgula.\n"
            "[GENEL_TREND_YORUM_BITIS]"
        )
        master_prompt_parts.append(
            "\n\n--- BÖLÜM: KATEGORİ TABLOSU YORUMU ---\n"
            "[KATEGORI_TABLO_YORUM_BASLANGIC]\n"
            "Yukarıda verilen 'Kategori Bazlı Yorum Sayıları'na dayanarak, yorumların konulara göre dağılımı hakkında TEK paragraflık bir yorum yaz. Hangi kategorilerin daha çok/az ilgi gördüğünü belirt ve bunun olası nedenleri veya anlamı hakkında uzun bir çıkarım yap.\n"
            "[KATEGORI_TABLO_YORUM_BITIS]"
        )
        
        for cat_data in category_counts_data:
            display_name = cat_data['name']
            csv_col_name = cat_data['csv_col']
            delimiter_key = f"KATEGORI_TREND_YORUM_{self._sanitize_for_delimiter(display_name)}"
            
            cat_pos_trend = "Bu kategori için pozitif duygu trend verisi bulunmamaktadır."
            cat_neg_trend = "Bu kategori için negatif duygu trend verisi bulunmamaktadır."

            if 'hesaplanan_tarih' in self.df_comments.columns and not self.df_comments['hesaplanan_tarih'].isna().all() and \
               csv_col_name in self.df_comments.columns:
                
                df_category_specific = self.df_comments[self.df_comments[csv_col_name].astype(bool)]

                if not df_category_specific.empty:
                    df_pos_agg_cat = df_category_specific[df_category_specific['duygu_tahmini'].str.lower() == 'positive'].groupby(pd.Grouper(key='hesaplanan_tarih', freq='D'))['duygu_skoru'].mean().reset_index()
                    df_neg_agg_cat = df_category_specific[df_category_specific['duygu_tahmini'].str.lower() == 'negative'].groupby(pd.Grouper(key='hesaplanan_tarih', freq='D'))['duygu_skoru'].mean().reset_index()
                    
                    cat_pos_trend = self._summarize_sentiment_trend(df_pos_agg_cat, "pozitif")
                    cat_neg_trend = self._summarize_sentiment_trend(df_neg_agg_cat, "negatif")

            master_prompt_parts.append(
                f"\n\n--- BÖLÜM: KATEGORİ ZAMAN SERİSİ YORUMU - {display_name} ---\n"
                f"'{display_name}' Kategorisi Verileri:\n"
                f"- Pozitif Duygu Trend Özeti: {cat_pos_trend}\n"
                f"- Negatif Duygu Trend Özeti: {cat_neg_trend}\n"
                f"[{delimiter_key}_BASLANGIC]\n"
                f"Yukarıdaki '{display_name}' kategorisine özel pozitif ve negatif duygu trend özetlerine dayanarak, bu kategorideki duygu skorlarının zaman içindeki değişimi hakkında TEK paragraflık uzun bir yorum yaz. Bu trendlerin '{display_name}' konusundaki kullanıcı algısı için ne anlama geldiğini analiz et.\n"
                f"[{delimiter_key}_BITIS]"
            )

        master_prompt_parts.append(
            "\n\n--- BÖLÜM: SONUÇ VE DEĞERLENDİRME ---\n"
            "[SONUC_DEGERLENDIRME_BASLANGIC]\n"
            f"Yukarıdaki tüm genel verilere ve sağlanan trend özetlerine dayanarak, '{self.keyword}' anahtar kelimesiyle yapılan YouTube yorum analizinin genel sonuçlarını ve değerlendirmesini  uzun bir metinle yap. Genel kullanıcı algısını, dikkat çeken kategorileri/duyguları, olası iyileştirme alanlarını veya gelecek stratejiler için önerileri belirt.\n"
            "[SONUC_DEGERLENDIRME_BITIS]"
        )
        master_prompt = "\n".join(master_prompt_parts)

        print("DEBUG: LLM çağrısı öncesi. Promptun bir kısmı:\n", master_prompt[:1000]) # Log a part of the prompt
        llm_full_response = self._ollama_gemma_chat(master_prompt)
        print("DEBUG: LLM çağrısı sonrası. Yanıtın başı:", llm_full_response[:200])

        if "LLM_YANIT_HATASI" in llm_full_response:
            print("LLM'den yanıt alınamadığı için rapor metinleri boş kalacak veya hata mesajı içerecektir.")
            ozet_text, pasta_yorum_text, genel_trend_yorum_text, kategori_tablo_yorum_text, sonuc_degerlendirme_text = ("LLM yanıt hatası.",) * 5
            kategori_trend_yorumları_parsed = {cat_data['name']: "LLM yanıt hatası." for cat_data in category_counts_data}
        else:
            print("DEBUG: LLM yanıtı parse ediliyor.")
            ozet_text = self._parse_llm_section(llm_full_response, "OZET")
            pasta_yorum_text = self._parse_llm_section(llm_full_response, "PASTA_YORUM")
            genel_trend_yorum_text = self._parse_llm_section(llm_full_response, "GENEL_TREND_YORUM")
            kategori_tablo_yorum_text = self._parse_llm_section(llm_full_response, "KATEGORI_TABLO_YORUM")
            sonuc_degerlendirme_text = self._parse_llm_section(llm_full_response, "SONUC_DEGERLENDIRME")
            kategori_trend_yorumları_parsed = {}
            for cat_data in category_counts_data:
                display_name = cat_data['name']
                delimiter_key = f"KATEGORI_TREND_YORUM_{self._sanitize_for_delimiter(display_name)}"
                kategori_trend_yorumları_parsed[display_name] = self._parse_llm_section(llm_full_response, delimiter_key)
            print("DEBUG: LLM yanıtı parse etme tamamlandı.")

        # --- RAPOR İÇERİĞİ ---
        try:
            document.add_paragraph(f'YouTube Yorum Analizi Raporu: "{self.keyword}"', style=title_style)
            document.add_paragraph(f"Analiz Tarihi: {datetime.now().strftime('%d-%m-%Y')}", style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph().paragraph_format.space_after = Pt(18)

            document.add_paragraph('1. ÖZET', style=heading1_style)
            document.add_paragraph(ozet_text, style=normal_style)
            document.add_page_break()

            document.add_paragraph('2. YÖNTEM', style=heading1_style)
            document.add_paragraph('2.1. Veri Toplama', style=heading2_style)
            data_collection_text = (
                f"Analiz için veriler, ‘{self.keyword}’ anahtar kelimesi kullanılarak YouTube platformunda yapılan arama ile başlatılmıştır. "
                f"Arama sonucunda hangi videonun başlığında veya açıklamasında '{self.keyword}' kelimesi geçiyorsa, bu videolara ait URL'ler "
                "CSV dosaysına kayıt edilmiştir. "
                f"Arama sonuçlarından, {video_count_placeholder} tane video URL'si kayıt edilmiştir. "
                f"Kayıt edilen URL'lerden  tek tek , yorum metni, tarih, kullanıcı adı ve diğer meta verilerle birlikte Web Scraping yöntemi kullanılarak toplam ' {total_comments} 'yorum  bir CSV dosyasına kaydedilmiştir. "
                "Bu CSV dosyası, yorumların işlenmesi ve analizi için temel veri kaynağı olarak kullanılmıştır."
            )
            document.add_paragraph(data_collection_text, style=normal_style)
            document.add_paragraph('2.2. Duygu Analizi ve Görselleştirme', style=heading2_style)
            sentiment_analysis_text = (
                "Toplanan ham yorum metinleri üzerinde duygu analizi gerçekleştirilmiştir. Bu süreçte, önceden eğitilmiş bir "
                "doğal dil işleme (NLP) modeli (BERT) kullanılmıştır. Her bir yorum için -1 (çok negatif) ile +1 (çok pozitif) "
                "arasında bir duygu skoru üretilmiştir. Bu skorlar daha sonra 'pozitif've 'negatif' "
                "olarak etiketlenmiştir."
                "Analiz sonuçları, genel duygu dağılımını gösteren bir pasta grafiği ve duygu skorlarının zaman içindeki "
                "değişimini gösteren bir zaman çizgisi grafiği ile görselleştirilmiştir."
            )
            document.add_paragraph(sentiment_analysis_text, style=normal_style)
            document.add_paragraph('2.3. Yorumların Kategorilere Göre Sınıflandırılması', style=heading2_style)
            categorization_text = (
                f"Yorumlar, içeriklerine göre önceden kullanıcı tarafından belirlenen {len(self.category_definitions)} ana kategoriye "
                f"ayrılmıştır. Belirlenen ana kategoriler şunlardır: {kategori_adlari_str_yontem}. "
                "Bu sınıflandırma işlemi, Büyük Dil Modeli (LLM) destekli bir yaklaşımla gerçekleştirilmiştir. LLM, her bir yorumun metnini analiz ederek "
                "hangi kategori veya kategorilere ait olduğuna dair bir etiket atamıştır. Bu yöntem, yorumların çoklu "
                "kategorilere atanabilmesine olanak tanımış ve daha esnek bir sınıflandırma sağlamıştır."
            )
            document.add_paragraph(categorization_text, style=normal_style)
            document.add_paragraph('2.4. Kategori Bazlı Analiz ve Görselleştirme', style=heading2_style)
            category_analysis_text = (
                "Her bir kategori için, o kategoriye ait yorumların sayısı ve bu yorumların genel duygu dağılımı incelenmiştir. "
                "Ayrıca, her bir kategorideki yorumların ortalama duygu skorlarının zaman içindeki değişimi de analiz edilerek "
                "ilgili kategoriye özel zaman çizgisi grafikleri oluşturulmuştur. Bu, belirli konular hakkındaki kullanıcı algısının "
                "zamanla nasıl evrildiğini anlamak için yapılmıştır."
            )
            document.add_paragraph(category_analysis_text, style=normal_style)
            document.add_paragraph('2.5. Raporlama', style=heading2_style)
            reporting_text = (
                "Yukarıda açıklanan tüm veri toplama, işleme, analiz ve görselleştirme adımlarının sonunda, elde edilen bulgular "
                "bu Word raporunda derlenmiştir. Rapor, analiz sürecini, kullanılan yöntemleri, temel bulguları ve "
                f"‘{self.keyword}’ anahtar kelimesi etrafındaki genel kullanıcı duyarlılığına dair sonuçları profesyonel ve anlaşılır "
                "bir şekilde sunmayı hedeflemektedir."
            )
            document.add_paragraph(reporting_text, style=normal_style)
            document.add_page_break()

            # --- 3. BULGULAR ---
            document.add_paragraph('3. BULGULAR', style=heading1_style)
            grafik_index = 0

            # 3.1. Genel Duygu Dağılımı
            document.add_paragraph('3.1. Yorumların Genel Duygu Dağılımı', style=heading2_style)
            pie_chart_buffer = self._create_modern_sentiment_pie_chart()
            if pie_chart_buffer.getbuffer().nbytes > 500: # Check if buffer has content
                grafik_index += 1
                document.add_picture(pie_chart_buffer, width=Inches(5.5)); document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_paragraph(f"Grafik {grafik_index}: Yorumların Genel Duygu Dağılımı ({total_comments} yorum üzerinden)", style=caption_style)
                document.add_paragraph(pasta_yorum_text, style=normal_style)
            else:
                document.add_paragraph("Genel duygu dağılımı grafiği için yeterli veri bulunmamaktadır.", style=normal_style)

            # 3.2. Zaman İçinde Genel Duygu Skoru Değişimi
            document.add_paragraph('3.2. Zaman İçinde Genel Duygu Skoru Değişimi', style=heading2_style)
            line_chart_buffer_general = self._create_modern_line_chart(self.df_comments, title=f'"{self.keyword}" Konulu Yorumlarda Genel Duygu Skoru Değişimi')
            if line_chart_buffer_general.getbuffer().nbytes > 1000: # Heuristic check for actual graph
                grafik_index += 1
                document.add_picture(line_chart_buffer_general, width=Inches(6.0)); document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                document.add_paragraph(f"Grafik {grafik_index}: Zaman İçinde Genel Duygu Skoru Değişimi", style=caption_style)
                document.add_paragraph(genel_trend_yorum_text, style=normal_style)
            else:
                document.add_paragraph("Genel duygu skoru zaman serisi grafiği için yeterli veri bulunmamaktadır.", style=normal_style)

            # 3.3. Yorumların Konulara Göre Dağılımı
            document.add_paragraph('3.3. Yorumların Konulara Göre Dağılımı', style=heading2_style)
            table_intro_text = (f"Analiz edilen toplam {total_comments} yorumun, önceden belirlenmiş kategorilere göre dağılımı "
                                f"aşağıdaki tabloda sunulmuştur. Her bir kategori için yorum sayısı ve bu sayının toplam yorumlara "
                                f"oranı yüzde cinsinden gösterilmektedir.")
            document.add_paragraph(table_intro_text, style=normal_style)

            if category_counts_data:
                table = document.add_table(rows=1, cols=3)
                table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False; table.columns[0].width = Inches(3.0); table.columns[1].width = Inches(1.5); table.columns[2].width = Inches(1.5)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Kategori Adı'; hdr_cells[1].text = 'Yorum Sayısı'; hdr_cells[2].text = 'Yüzdesi (%)'
                header_bg_color = "D9EAD3" 
                for cell in hdr_cells:
                    cell.paragraphs[0].runs[0].font.bold = True; cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._set_cell_shading(cell, header_bg_color)

                for item in sorted(category_counts_data, key=lambda x: x['count'], reverse=True):
                    row_cells = table.add_row().cells
                    row_cells[0].text = item['name']
                    row_cells[1].text = str(item['count']); row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    percentage = (item['count'] / total_comments * 100) if total_comments > 0 else 0
                    row_cells[2].text = f"{percentage:.1f}%"; row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                total_row_cells = table.add_row().cells
                total_row_cells[0].text = 'Toplam'; total_row_cells[0].paragraphs[0].runs[0].font.bold = True
                total_row_cells[1].text = str(total_comments); total_row_cells[1].paragraphs[0].runs[0].font.bold = True; total_row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                total_row_cells[2].text = "100.0%" if total_comments > 0 else "0.0%"; total_row_cells[2].paragraphs[0].runs[0].font.bold = True; total_row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                total_row_bg_color = "F3F3F3" 
                for cell in total_row_cells: self._set_cell_shading(cell, total_row_bg_color)
                
                document.add_paragraph(f"Tablo 1: Yorumların Konulara Göre Dağılımı", style=caption_style)
                document.add_paragraph(kategori_tablo_yorum_text, style=normal_style)
            else:
                document.add_paragraph("Kategorilere göre yorum dağılım tablosu için veri bulunmamaktadır.", style=normal_style)
            
            # 3.4. Kategori Bazlı Duygu Skorlarının Zaman İçindeki Değişimi
            document.add_paragraph('3.4. Kategori Bazlı Duygu Skorlarının Zaman İçindeki Değişimi', style=heading2_style)
            for cat_data in category_counts_data:
                display_name = cat_data['name']
                csv_col_name = cat_data['csv_col']
                
                document.add_paragraph(f"Konu: {display_name}", style=heading2_style)
                # Ensure the column exists before filtering
                if csv_col_name not in self.df_comments.columns:
                    no_col_msg = f"'{display_name}' kategorisi için veri sütunu ('{csv_col_name}') DataFrame'de bulunamadı."
                    document.add_paragraph(no_col_msg, style=normal_style)
                    document.add_paragraph(kategori_trend_yorumları_parsed.get(display_name, ""), style=normal_style)
                    continue

                df_category_specific = self.df_comments[self.df_comments[csv_col_name].astype(bool)]
                
                if df_category_specific.empty:
                    no_data_msg = f"'{display_name}' kategorisine ait yorum bulunmadığından bu kategori için zaman serisi grafiği oluşturulamamıştır."
                    document.add_paragraph(no_data_msg, style=normal_style)
                    document.add_paragraph(kategori_trend_yorumları_parsed.get(display_name, ""), style=normal_style)
                    continue
                    
                category_line_chart_buffer = self._create_modern_line_chart(df_category_specific, title=f"{display_name} - Duygu Skoru Değişimi")
                if category_line_chart_buffer.getbuffer().nbytes > 1000: # Heuristic
                    grafik_index += 1
                    document.add_picture(category_line_chart_buffer, width=Inches(6.0)); document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    document.add_paragraph(f"Grafik {grafik_index}: {display_name} - Zaman İçinde Duygu Skoru Değişimi", style=caption_style)
                    document.add_paragraph(kategori_trend_yorumları_parsed.get(display_name, "Bu kategori için LLM yorumu bulunamadı."), style=normal_style)
                else:
                    no_graph_msg = f"'{display_name}' kategorisi için zaman serisi grafiği oluşturulamadı (yetersiz veya uygun olmayan veri)."
                    document.add_paragraph(no_graph_msg, style=normal_style)
                    document.add_paragraph(kategori_trend_yorumları_parsed.get(display_name, ""), style=normal_style)
            document.add_page_break()

            document.add_paragraph('4. SONUÇ ve DEĞERLENDİRME', style=heading1_style)
            document.add_paragraph(sonuc_degerlendirme_text, style=normal_style)

            document.save(output_filename)
            print(f"Profesyonel rapor '{output_filename}' başarıyla oluşturuldu.")

        except Exception as e_report_gen:
            print(f"RAPOR OLUŞTURMA SIRASINDA BEKLENMEDİK HATA: {e_report_gen}")
            import traceback
            traceback.print_exc()
            try:
                if document:
                    error_paragraph = document.add_paragraph(f"Rapor içeriği oluşturulurken ciddi bir hata oluştu: {e_report_gen}\n{traceback.format_exc()}", style=normal_style)
                    document.save(f"HATALI_{output_filename}")
                    print(f"Hatalı rapor 'HATALI_{output_filename}' olarak kaydedildi.")
            except Exception as e_save_error:
                print(f"Hatalı raporu kaydederken ek hata: {e_save_error}")

